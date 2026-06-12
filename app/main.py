import streamlit as st
import os
import sys
import io
import shutil
from pathlib import Path
from datetime import datetime

APP_DIR = Path(__file__).parent
sys.path.insert(0, str(APP_DIR))

from rag import RAG
from model import gpt_analysis, gpt_with_retry, chat
from prompts import (
    get_topic_selection_prompt,
    get_modeling_prompt,
    get_coding_prompt,
    get_figure_prompt,
    get_paper_writing_prompt,
    get_polish_prompt,
    get_pressure_test_prompt,
    get_grill_me_prompt
)
from workflow_logger import WorkflowLogger
from skills_bridge import SCIPILOT_DIR, GPT_ACADEMIC_DIR, SKILLS_DIR
from skills_runner  import run_pressure_test, run_grill_me, run_code_check, get_skill_comparison
from prd_generator  import generate_prd, generate_claude_md, export_prd_file
from quota_monitor  import QuotaMonitor
from app.inbox_watcher import (
    ensure_inbox_dirs, get_inbox_status, get_new_files,
    mark_as_processed, read_data_file
)
from app.webai_bridge import (
    export_context_package, import_webai_response, list_export_files
)
from app.memory_logger import MemoryLogger
from app.utils import load_tabular_file, list_data_files
from app.session_persistence import (
    save_session, load_session, list_sessions,
    copy_uploaded_file, get_upload_dir, get_session_id_from_url,
    save_latest_session_id, get_latest_session_id
)
from app.context_restorer import (
    rebuild_rag_from_session, build_context_summary, needs_rag_rebuild
)

st.set_page_config(page_title="APMCM 数学建模 Agent", page_icon="🎓", layout="wide")


@st.cache_resource
def init_rag():
    return RAG(
        problems_dir="dataset/problems",
        papers_dir="dataset/papers",
        references_dir="dataset/references"
    )


if "rag" not in st.session_state:
    with st.spinner("正在初始化 RAG 引擎，首次加载需要向量化全部 PDF (~1分钟)..."):
        st.session_state.rag = init_rag()
if "logger" not in st.session_state:
    st.session_state.logger = WorkflowLogger()
if "phase" not in st.session_state:
    st.session_state.phase = "input"
if "topics" not in st.session_state:
    st.session_state.topics = []
if "selected_topic" not in st.session_state:
    st.session_state.selected_topic = None
if "modeling_plan" not in st.session_state:
    st.session_state.modeling_plan = None
if "coding_result" not in st.session_state:
    st.session_state.coding_result = None
if "figure_descriptions" not in st.session_state:
    st.session_state.figure_descriptions = None
if "paper_draft" not in st.session_state:
    st.session_state.paper_draft = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "_init_done" not in st.session_state:
    st.session_state._init_done = True

    if "session_id" not in st.session_state:
        recovered_id = (
            get_session_id_from_url(dict(st.query_params))
            or get_latest_session_id()
        )
        if recovered_id:
            st.session_state.session_id = recovered_id
        else:
            st.session_state.session_id = datetime.now().strftime("%Y%m%d_%H%M%S")

    try:
        st.query_params["session"] = st.session_state.session_id
    except Exception:
        pass
    save_latest_session_id(st.session_state.session_id)

    saved = load_session(st.session_state.session_id)
    if saved:
        for k, v in saved.items():
            if k not in ("_init_done",):
                st.session_state[k] = v
        st.session_state.state_restored = True
    else:
        st.session_state.state_restored = False

for k, v in {
    "uploaded_file_paths": [],
    "rag_rebuilt": False,
    "context_summary": "",
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ══ DEBUG PANEL ══
with st.sidebar:
    with st.expander("🔍 DEBUG", expanded=False):
        import json as _json
        sid = st.session_state.get("session_id", "?")
        st.write(f"session_id: `{sid}`")
        st.write(f"state_restored: `{st.session_state.get('state_restored')}`")
        st.write(f"rag_rebuilt: `{st.session_state.get('rag_rebuilt')}`")
        json_path = Path(__file__).resolve().parent.parent / "workspace" / sid / "session_state.json"
        st.write(f"json存在: `{json_path.exists()}`")
        if json_path.exists():
            d = _json.loads(json_path.read_text(encoding="utf-8"))
            st.write(f"phase: `{d.get('phase')}`")
            st.write(f"chat_history: {len(d.get('chat_history', []))} 条")
            st.write(f"uploaded_file_paths: {len(d.get('uploaded_file_paths', []))} 个")
        up_dir = Path(__file__).resolve().parent.parent / "workspace" / sid / "uploads"
        st.write(f"uploads目录: `{up_dir.exists()}` 文件数: {len(list(up_dir.iterdir())) if up_dir.exists() else 0}")

if "quota_monitor" not in st.session_state:
    st.session_state.quota_monitor = QuotaMonitor(platform="Claude Sonnet")
if "completed_stages" not in st.session_state:
    st.session_state.completed_stages = []
if "prd_draft" not in st.session_state:
    st.session_state.prd_draft = ""
if "prd_final" not in st.session_state:
    st.session_state.prd_final = ""
if "grill_rounds" not in st.session_state:
    st.session_state.grill_rounds = 0
if "pressure_report" not in st.session_state:
    st.session_state.pressure_report = ""
if "paper_sections" not in st.session_state:
    st.session_state.paper_sections = {}
if "reference_loaded" not in st.session_state:
    st.session_state.reference_loaded = False
if "memory_logger" not in st.session_state:
    st.session_state.memory_logger = None
if "webai_import_filename" not in st.session_state:
    st.session_state.webai_import_filename = ""
if "webai_imported_content" not in st.session_state:
    st.session_state.webai_imported_content = ""
if "inbox_last_scan" not in st.session_state:
    st.session_state.inbox_last_scan = None
if "loaded_data_files" not in st.session_state:
    st.session_state.loaded_data_files = {}

if st.session_state.memory_logger is None:
    st.session_state.memory_logger = MemoryLogger(st.session_state.session_id)

ensure_inbox_dirs()

logger = st.session_state.logger
rag = st.session_state.rag

PHASE_DOWNSTREAM = {
    "input":        ["topic_sims","topic_scores","topic_recommendation","modeling_plan","coding_result","paper_draft","pressure_report","prd_draft","prd_final","figure_descriptions","polished_paper","paper_sections","selected_topic","selected_topic_idx","selected_sims","completed_stages","grill_rounds","pressure_test_result","grill_result"],
    "topic_selection":["modeling_plan","coding_result","paper_draft","pressure_report","prd_draft","prd_final","figure_descriptions","polished_paper","paper_sections","selected_topic","selected_topic_idx","selected_sims","completed_stages","grill_rounds","pressure_test_result","grill_result"],
    "modeling":      ["pressure_report","prd_draft","prd_final","coding_result","paper_draft","figure_descriptions","polished_paper","paper_sections","grill_result"],
    "coding":        ["paper_draft","figure_descriptions","polished_paper","paper_sections"],
    "figure":        ["paper_draft","polished_paper"],
    "paper":         ["polished_paper"],
}

PHASE_FILES = {
    "modeling":      ["PRD.md", "CLAUDE.md", "prepare_claude/"],  # modeling stage files in workspace/
    "coding":        [],  # handled by folder move
    "figure":        [],  # picture/ folder
    "paper":         [],  # writing/ folder
}

def rubbish_downstream(phase):
    """归档下游结果：移入 workspace/rubbish/{phase}_{timestamp}/，并清除 session_state"""
    from datetime import datetime
    rb = Path(__file__).resolve().parent.parent / "workspace" / "rubbish"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    rb_dir = rb / f"{phase}_{ts}"
    rb_dir.mkdir(parents=True, exist_ok=True)

    ws = Path(__file__).resolve().parent.parent / "workspace"
    n = st.session_state.get("selected_topic_idx", 0) + 1

    # 移动 workspace/ 下的阶段文件
    for fname in PHASE_FILES.get(phase, []):
        src = ws / fname
        if src.exists():
            if src.is_dir():
                shutil.move(str(src), str(rb_dir / fname))
            else:
                shutil.move(str(src), str(rb_dir / fname))

    # 移动文件夹（按阶段）
    if phase in ("modeling", "topic_selection"):
        for d in ["coding", "writing", "picture"]:
            src = ws / d
            if src.exists():
                shutil.move(str(src), str(rb_dir / d))
                # 重建空目录
                src.mkdir(parents=True, exist_ok=True)
                (src / "upload").mkdir(parents=True, exist_ok=True)
        (ws / "coding").mkdir(parents=True, exist_ok=True)
        (ws / "writing").mkdir(parents=True, exist_ok=True)
        (ws / "picture").mkdir(parents=True, exist_ok=True)
        (ws / "upload").mkdir(parents=True, exist_ok=True)
    elif phase == "coding":
        for d in [f"coding/{n}", f"writing/{n}", "picture"]:
            src = ws / d
            if src.exists():
                shutil.move(str(src), str(rb_dir / d))
                src.mkdir(parents=True, exist_ok=True)
    elif phase == "figure":
        for d in [f"writing/{n}"]:
            src = ws / d
            if src.exists():
                shutil.move(str(src), str(rb_dir / d))
                src.mkdir(parents=True, exist_ok=True)

    # 清除 session_state 下游键
    for key in PHASE_DOWNSTREAM.get(phase, []):
        st.session_state.pop(key, None)

    if any(rb_dir.iterdir()):
        st.toast(f"📦 旧结果已归档到 rubbish/{phase}_{ts[:13]}", icon="📦")

def autosave():
    """保存当前 session 状态到磁盘（带错误保护）"""
    ok = save_session(st.session_state.session_id, dict(st.session_state))
    save_latest_session_id(st.session_state.session_id)
    return ok

if (st.session_state.state_restored
        and not st.session_state.rag_rebuilt
        and needs_rag_rebuild(st.session_state.get("rag"))):
    with st.spinner("🔄 检测到重启，正在重新扫描文件恢复上下文..."):
        report = rebuild_rag_from_session(
            st.session_state.session_id,
            st.session_state.rag
        )
        st.session_state.rag_rebuilt = True
        st.session_state.context_summary = build_context_summary(dict(st.session_state))
        if report["total"] > 0:
            st.toast(f"✅ 已恢复 {report['total']} 个文件的知识库索引", icon="📚")

# === 全局侧边栏：进度 + 额度监控 + 提问入口 ===
with st.sidebar:
    with st.expander("💾 会话管理", expanded=False):
        st.caption(f"当前 Session: `{st.session_state.session_id}`")
        phase_label = st.session_state.get("phase", "input")
        st.caption(f"当前阶段: **{phase_label}**")
        if st.session_state.get("state_restored"):
            st.success("✅ 已从磁盘恢复上次进度")
        if st.button("💾 立即保存进度", key="btn_manual_save"):
            save_session(st.session_state.session_id, dict(st.session_state))
            st.toast("✅ 进度已保存", icon="💾")
        st.divider()
        st.caption("历史会话（点击可恢复）：")
        sessions = list_sessions()
        if sessions:
            for s in sessions[:5]:
                label = f"[{s['phase']}] {s['session_id']} · {s['saved_at'][:16]}"
                if st.button(label, key=f"restore_{s['session_id']}"):
                    target_id = s["session_id"]
                    saved = load_session(target_id)
                    if saved:
                        keys_to_clear = [k for k in st.session_state.keys() if k != "_init_done"]
                        for key in keys_to_clear:
                            del st.session_state[key]
                        for k, v in saved.items():
                            st.session_state[k] = v
                        st.session_state.session_id = target_id
                        st.session_state.state_restored = True
                        st.session_state.rag_rebuilt = False
                        st.session_state.context_summary = ""
                        save_latest_session_id(target_id)
                        try:
                            st.query_params["session"] = target_id
                        except Exception:
                            pass
                        st.rerun()
                    else:
                        st.error(f"无法加载会话 {target_id}，session_state.json 不存在")
        else:
            st.caption("（暂无历史会话）")
        st.divider()
        st.caption("本次已上传文件：")
        upload_dir = get_upload_dir(st.session_state.session_id)
        all_files = list(upload_dir.glob("*")) if upload_dir.exists() else []
        if all_files:
            for f in all_files:
                st.text(f"📄 {f.name}")
        else:
            st.caption("（暂无）")

    st.divider()

    st.header("📊 项目仪表盘")

    # 进度追踪
    all_stages = ["选题确认", "建模方法确认", "压力测试", "PRD生成",
                  "需求对齐", "代码生成", "图表生成", "论文撰写", "论文润色"]
    completed  = st.session_state.completed_stages
    progress   = len(completed) / len(all_stages)

    st.progress(progress, text=f"整体进度 {len(completed)}/{len(all_stages)}")
    for s in all_stages:
        icon = "✅" if s in completed else ("🔄" if s == st.session_state.get("phase", "") else "⏳")
        st.caption(f"{icon} {s}")

    st.divider()

    # 额度监控
    qm = st.session_state.quota_monitor
    st.subheader("💰 Token 额度")

    platform_choice = st.selectbox(
        "当前平台",
        ["Claude Sonnet", "Claude Opus 4.8", "ChatGPT-4o", "ChatGPT-5.5"],
        key="platform_select"
    )
    if platform_choice != qm.platform:
        st.session_state.quota_monitor = QuotaMonitor(platform=platform_choice)
        qm = st.session_state.quota_monitor

    st.code(qm.get_status_text(), language=None)

    # 额度预警
    should_alert, threshold = qm.should_alert()
    if should_alert:
        workspace = Path(__file__).resolve().parent.parent / "workspace" / st.session_state.get("session_id", "default")
        handoff   = qm.generate_handoff_doc(
            session_id     = st.session_state.get("session_id", "default"),
            current_stage  = st.session_state.get("phase", "未知"),
            completed_stages = completed,
            workspace_path = str(workspace.absolute()),
            prd_summary    = st.session_state.get("prd_final", "")[:500],
        )
        st.warning(f"⚠️ 额度已用 {threshold*100:.0f}%，建议准备切换平台")
        with st.expander("📋 查看任务交接文档"):
            st.markdown(handoff)
            buf = io.BytesIO(handoff.encode("utf-8"))
            st.download_button("⬇️ 下载交接文档", buf, "任务交接.md", "text/markdown")

    st.divider()

    # 文件收件箱
    with st.expander("📥 文件收件箱", expanded=False):
        st.caption("将文件直接拖入对应目录，agent 启动时自动加载")
        inbox_path = Path(__file__).resolve().parent.parent / "inbox"
        st.code(str(inbox_path), language=None)
        if st.button("🔄 扫描新文件", key="btn_scan_inbox"):
            status = get_inbox_status()
            for subdir, info in status.items():
                label = {"problems":"赛题","papers":"论文","references":"参考文献",
                         "knowledge":"知识库","web_ai":"网页AI回复"}.get(subdir, subdir)
                new_badge = f"  🆕 {info['new']} 个新文件" if info["new"] > 0 else ""
                st.write(f"**{label}**: {info['total']} 个文件{new_badge}")
            total_loaded = 0
            _rag = st.session_state.get("rag")
            if _rag:
                for sub in ["problems","papers","references","knowledge"]:
                    new_files = get_new_files(sub)
                    if new_files:
                        n = _rag.add_inbox_files(sub, new_files)
                        mark_as_processed(new_files)
                        total_loaded += n
                if total_loaded > 0:
                    st.success(f"✅ 已加载 {total_loaded} 个新文件到知识库")
        st.caption("📂 子目录：problems / papers / references / knowledge / web_ai / data")

        st.divider()
        st.caption("📊 数据文件（CSV / Excel）")
        st.caption("将数据文件放入 inbox/data/ 目录后点击加载：")

        data_path = Path(__file__).resolve().parent.parent / "inbox" / "data"
        st.code(str(data_path), language=None)

        data_files = list_data_files(data_path)
        if data_files:
            for f in data_files:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.text(f"📄 {f['name']} ({f['size_kb']} KB)")
                with col2:
                    btn_key = f"load_data_{f['name']}"
                    if st.button("加载", key=btn_key):
                        summary = read_data_file(f["name"])
                        st.session_state.loaded_data_files[f["name"]] = summary
                        st.session_state.memory_logger.log_system_event(
                            "加载数据文件", f["name"]
                        )
                        st.success(f"✅ 已加载 {f['name']}")
                        st.text_area("预览", value=summary[:300]+"...", height=100, disabled=True)
        else:
            st.caption("（inbox/data/ 目录下暂无数据文件）")

        if st.session_state.get("loaded_data_files"):
            names = list(st.session_state.loaded_data_files.keys())
            st.info(f'💡 已加载 {len(names)} 个数据文件：{", ".join(names)}\n可在提问时说"分析数据"自动引用')

    st.divider()

    # AI 网页版协作
    with st.expander("🔗 AI 网页版协作", expanded=False):
        st.caption("将当前阶段上下文导出，上传给网页版 AI；或读取网页版 AI 的回复")
        if st.button("📤 导出上下文给网页版 AI", key="btn_export_webai"):
            workspace_root = Path(__file__).resolve().parent.parent / "workspace"
            export_path = export_context_package(
                phase=st.session_state.get("phase", "input"),
                session_id=st.session_state.session_id,
                context=dict(st.session_state),
                workspace_root=workspace_root
            )
            st.success("✅ 已导出")
            st.code(str(export_path), language=None)
            st.caption("请到上述路径找到文件，上传给网页版 AI")
        st.divider()
        st.caption("将网页版 AI 的回复保存到 inbox/web_ai/ 目录后，在此输入文件名：")
        filename_input = st.text_input(
            "回复文件名（如 chatgpt_reply.md）",
            key="webai_filename_input",
            placeholder="chatgpt_reply.md"
        )
        if st.button("📥 读取网页版 AI 回复", key="btn_import_webai"):
            if filename_input.strip():
                content = import_webai_response(filename_input.strip())
                st.session_state.webai_imported_content = content
                if content.startswith("❌"):
                    st.error(content)
                else:
                    st.success(f"✅ 已读取 {len(content)} 字符")
                    st.text_area("回复内容预览", value=content[:500]+"...", height=150, disabled=True)
            else:
                st.warning("请先输入文件名")
        if st.session_state.get("webai_imported_content"):
            st.info("💡 网页版 AI 回复已就绪，可在对话框中输入\"使用网页AI方案\"来应用")
        st.divider()
        st.caption("本次 session 已导出的文件：")
        workspace_root_exports = Path(__file__).resolve().parent.parent / "workspace"
        exports = list_export_files(st.session_state.session_id, workspace_root_exports)
        if exports:
            for e in exports[-5:]:
                st.text(f"📄 {e['filename']}")
        else:
            st.caption("（暂无导出记录）")

    st.divider()

    # Skill 对比查询
    with st.expander("🛠️ Skills 说明"):
        st.markdown(get_skill_comparison())

    st.divider()

    # 用户随时提问
    st.subheader("💬 随时提问")
    user_question = st.text_input("对当前项目有什么疑问？", key="sidebar_question", placeholder="例如：为什么选这个模型？")
    if user_question and st.button("提问", key="sidebar_ask"):
        st.session_state.memory_logger.log_message("user", user_question)
        ctx = f"""项目进度：{completed}
当前阶段：{st.session_state.get('phase', '未知')}
建模方案摘要：{(st.session_state.get('modeling_plan') or '')[:400]}
PRD摘要：{(st.session_state.get('prd_final') or st.session_state.get('prd_draft') or '')[:400]}"""
        from model import gpt_with_retry
        answer = gpt_with_retry(f"用户问题：{user_question}\n\n项目背景：{ctx}", max_tokens=500)
        qm.record("用户提问", qm.estimate_tokens(answer))
        st.session_state.memory_logger.log_message("assistant", answer)
        st.info(answer)

st.title("🎓 APMCM 数学建模比赛 Agent")
st.caption("基于 RAG + LLM + 多Skill 协作的数学建模助手")

# ============ Phase 1: Input ============
if st.session_state.phase == "input":
    st.header("📝 上传赛题 PDF")

    UPLOAD_ROOT = Path(__file__).resolve().parent.parent / "workspace" / "upload"
    for d in [UPLOAD_ROOT / "1", UPLOAD_ROOT / "2", UPLOAD_ROOT / "3"]:
        d.mkdir(parents=True, exist_ok=True)

    def _extract_pdf(file_bytes):
        from pypdf import PdfReader
        from io import BytesIO
        txt = ""
        pdf = PdfReader(BytesIO(file_bytes))
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                txt += extracted + "\n"
        return txt.strip()

    def _read_folder_topic(folder_num):
        """从 workspace/upload/{folder_num}/ 读取第一个 PDF 文件"""
        folder = UPLOAD_ROOT / str(folder_num)
        if folder.exists():
            pdfs = list(folder.glob("*.pdf"))
            if pdfs:
                return pdfs[0]
        return None

    # 扫描各文件夹中的已有文件
    st.markdown(f"文件路径: `{UPLOAD_ROOT.absolute()}`")
    st.caption("将赛题 PDF 放入对应子文件夹（1/2/3），或通过下方上传区域上传。要删除选题，从对应文件夹中删除 PDF 文件即可。")

    col1, col2, col3 = st.columns(3)
    uploaded_texts = [None, None, None]
    uploaded_names = [None, None, None]
    topic_files = [None, None, None]

    for i, col in enumerate([col1, col2, col3]):
        folder_num = str(i + 1)
        folder = UPLOAD_ROOT / folder_num
        with col:
            st.subheader(f"选题 {folder_num}")

            # 显示已有文件
            existing_pdf = _read_folder_topic(folder_num)
            if existing_pdf:
                topic_files[i] = existing_pdf
                st.info(f"📄 {existing_pdf.name}")
                if st.button("🗑️ 删除", key=f"del_topic_{i}"):
                    existing_pdf.unlink()
                    st.rerun()

            # 上传区域
            uploaded = st.file_uploader(
                f"上传PDF", type=["pdf"], key=f"upload_{i}",
                help=f"上传赛题{folder_num}，会保存到 upload/{folder_num}/"
            )
            if uploaded:
                dest = folder / uploaded.name
                dest.write_bytes(uploaded.getvalue())
                with st.spinner(f"解析选题{folder_num}..."):
                    topic_files[i] = dest
                    st.session_state[f"extracted_{i}"] = _extract_pdf(dest.read_bytes())
                st.session_state[f"uploaded_name_{i}"] = uploaded.name
                autosave()
                st.success(f"已保存: {uploaded.name}")
                st.rerun()

    # 从已有文件读取文本
    for i in range(3):
        if topic_files[i]:
            pdf_path = topic_files[i]
            st.session_state[f"extracted_{i}"] = _extract_pdf(pdf_path.read_bytes())
            st.session_state[f"uploaded_name_{i}"] = pdf_path.name
            uploaded_texts[i] = st.session_state[f"extracted_{i}"]
            uploaded_names[i] = pdf_path.name

    if st.button("🚀 开始分析", type="primary", use_container_width=True):
        topics = []
        for i in range(3):
            t = uploaded_texts[i] or st.session_state.get(f"extracted_{i}")
            if t:
                topics.append(t)
                uploaded_names[i] = st.session_state.get(f"uploaded_name_{i}", f"选题{i+1}")
        if not topics:
            st.error("请至少上传一个赛题 PDF")
        else:
            # 归档旧结果，从当前文件重新开始
            rubbish_downstream("topic_selection")
            st.session_state.topics = topics
            logger.log_user_input(f"用户上传了{len(topics)}个赛题PDF: {[n for n in uploaded_names if n]}")
            st.session_state.phase = "topic_selection"
            st.session_state.memory_logger.new_stage(st.session_state.phase)
            st.session_state.memory_logger.log_system_event(
                "进入阶段: topic_selection",
                f"session_id={st.session_state.session_id}"
            )
            autosave()
            st.rerun()

# ============ Phase 2: Topic Selection ============
elif st.session_state.phase == "topic_selection":
    st.header("📊 选题分析")

    if st.button("⬅️ 返回上传赛题", key="back_to_input"):
        st.session_state.phase = "input"
        st.session_state.memory_logger.new_stage("input")
        st.session_state.memory_logger.log_system_event("返回阶段: input", "用户从选题分析返回")
        autosave()
        st.rerun()
    # → 下一阶段（选题→建模）
    has_modeling = st.session_state.get("modeling_plan") is not None
    if st.button("→ 建模方案", key="next_to_modeling", disabled=not has_modeling,
                 help="需先生成建模方案" if not has_modeling else "进入建模方案阶段"):
        st.session_state.phase = "modeling"
        st.session_state.memory_logger.new_stage("modeling")
        st.session_state.memory_logger.log_system_event("进入阶段: modeling")
        autosave()
        st.rerun()

    topics = st.session_state.topics
    if "topic_sims" not in st.session_state:
        with st.spinner("正在检索历史数据,分析各选题..."):
            sims_list = []
            scores = []
            for i, topic in enumerate(topics):
                with st.status(f"分析选题 {i+1}..."):
                    score, sims = rag.topic_coverage_score(topic, topk=5)
                    scores.append(score)
                    sims_list.append(sims)
                    st.write(f"选题{i+1} 匹配分数: {score}")
                    st.write(f"  - 相似历史题: {len(sims['sim_questions'])} 道")
                    st.write(f"  - 相关论文: {len(sims['sim_papers'])} 篇")
                    st.write(f"  - 参考文献: {len(sims['sim_refs'])} 篇")

            prompt = get_topic_selection_prompt(topics, sims_list)
            recommendation = gpt_with_retry(prompt)
            st.session_state.topic_recommendation = recommendation
            st.session_state.topic_sims = sims_list
            st.session_state.topic_scores = scores
            logger.log_topic_selection(topics, scores, recommendation)

    st.markdown("### 🤖 Agent 选题推荐")
    st.markdown(st.session_state.topic_recommendation)

    st.divider()
    st.markdown("### 📊 数据库匹配分数")
    for i, (topic, score, sims) in enumerate(zip(topics, st.session_state.topic_scores, st.session_state.topic_sims)):
        st.metric(f"选题{i+1}", f"{score} 分",
                  f"相似题:{len(sims['sim_questions'])} | 论文:{len(sims['sim_papers'])} | 参考:{len(sims['sim_refs'])}")

    col_a, col_b = st.columns(2)
    with col_a:
        selected_idx = st.selectbox("选择最终选题", range(len(topics)),
                                    format_func=lambda x: f"选题{x+1}")
    with col_b:
        if st.button("✅ 确认选题,开始建模", type="primary"):
            st.session_state.selected_topic = topics[selected_idx]
            st.session_state.selected_topic_idx = selected_idx
            st.session_state.selected_sims = st.session_state.topic_sims[selected_idx]
            logger.log_user_feedback("topic_selection", f"用户选择选题{selected_idx+1}")
            st.session_state.phase = "modeling"
            st.session_state.memory_logger.new_stage(st.session_state.phase)
            st.session_state.memory_logger.log_system_event(
                "进入阶段: modeling",
                f"session_id={st.session_state.session_id}"
            )
            autosave()
            st.rerun()

# ============ Phase 3: Modeling ============
elif st.session_state.phase == "modeling":
    st.header("🔬 数学建模方案")

    if st.button("⬅️ 返回选题分析", key="back_to_topic"):
        st.session_state.phase = "topic_selection"
        st.session_state.memory_logger.new_stage("topic_selection")
        st.session_state.memory_logger.log_system_event("返回阶段: topic_selection")
        autosave()
        st.rerun()
    has_coding = st.session_state.get("coding_result") is not None
    if st.button("→ 代码生成", key="next_to_coding", disabled=not has_coding,
                 help="需先生成代码" if not has_coding else "进入代码生成阶段"):
        st.session_state.phase = "coding"
        st.session_state.memory_logger.new_stage("coding")
        st.session_state.memory_logger.log_system_event("进入阶段: coding")
        autosave()
        st.rerun()

    selected_topic = st.session_state.selected_topic
    selected_sims = st.session_state.selected_sims

    # 显示 inbox/data/ 中可用的数据文件
    data_path = Path(__file__).resolve().parent.parent / "inbox" / "data"
    available_data = list_data_files(data_path)
    if available_data:
        data_names = [f["name"] for f in available_data]
        selected_data = st.multiselect(
            "选择需要使用的数据文件（将加入建模上下文）",
            options=data_names,
            default=data_names[:1] if data_names else [],
            help="选中的文件内容会注入到建模方案生成 Prompt 中"
        )
    else:
        selected_data = []

    if "modeling_plan" not in st.session_state or st.session_state.modeling_plan is None:
        with st.spinner("正在生成建模方案..."):
            data_context = ""
            if selected_data:
                for name in selected_data:
                    if name in st.session_state.get("loaded_data_files", {}):
                        data_context += f"\n\n### 数据文件: {name}\n{st.session_state.loaded_data_files[name]}"
                    else:
                        from app.inbox_watcher import read_data_file
                        summary = read_data_file(name)
                        st.session_state.loaded_data_files[name] = summary
                        data_context += f"\n\n### 数据文件: {name}\n{summary}"
            approach = st.session_state.get("user_approach", "")
            prompt = get_modeling_prompt(selected_topic, selected_sims, approach)
            if data_context:
                prompt += f"\n\n## 可用数据文件\n{data_context}"
            refs = st.session_state.pop("_modeling_refs", "")
            if refs:
                prompt += f"\n\n## 外部参考文件\n{refs}\n\n请综合以上参考内容重新优化建模方案。"
            modeling_plan = gpt_with_retry(prompt, max_tokens=6000)
            st.session_state.modeling_plan = modeling_plan
            logger.log_modeling(modeling_plan)
            autosave()

    st.markdown(st.session_state.modeling_plan)
    if st.button("💾 保存建模方案到 prepare_claude/", key="save_modeling_plan"):
        prep_dir = Path(__file__).resolve().parent.parent / "workspace" / "prepare_claude"
        prep_dir.mkdir(parents=True, exist_ok=True)
        (prep_dir / "建模方案.md").write_text(st.session_state.modeling_plan, encoding="utf-8")
        st.success(f"已保存到 `{prep_dir.absolute() / '建模方案.md'}`")

    # === 步骤A：压力测试 ===
    if st.session_state.get("modeling_plan") and not st.session_state.get("pressure_report"):
        if st.button("🔬 运行压力测试", key="run_pressure"):
            with st.spinner("正在运行压力测试..."):
                report = run_pressure_test(
                    question      = st.session_state.selected_topic,
                    modeling_plan = st.session_state.modeling_plan,
                )
                st.session_state.pressure_report = report
                qm.record("压力测试", qm.estimate_tokens(report))
                if "压力测试" not in st.session_state.completed_stages:
                    st.session_state.completed_stages.append("压力测试")
                autosave()

    if st.session_state.get("pressure_report"):
        with st.expander("📋 压力测试报告", expanded=True):
            st.markdown(st.session_state.pressure_report)
        if st.button("💾 保存报告到 prepare_claude/", key="save_pressure"):
            prep_dir = Path(__file__).resolve().parent.parent / "workspace" / "prepare_claude"
            prep_dir.mkdir(parents=True, exist_ok=True)
            (prep_dir / "压力测试报告.md").write_text(st.session_state.pressure_report, encoding="utf-8")
            st.success(f"已保存到 `{prep_dir.absolute() / '压力测试报告.md'}`")

    # === 步骤B：生成并保存 PRD 到 workspace/ 文件夹 ===
    if st.session_state.get("modeling_plan") and not st.session_state.get("prd_draft"):
        if st.button("📄 生成 PRD 并保存到 workspace/", key="gen_prd"):
            workspace_dir = Path(__file__).resolve().parent.parent / "workspace"
            with st.spinner("正在生成 PRD..."):
                prd = generate_prd(
                    question       = st.session_state.selected_topic,
                    modeling_plan  = st.session_state.modeling_plan,
                    pressure_report= st.session_state.get("pressure_report", ""),
                    align_context  = {},
                    session_id     = st.session_state.session_id,
                )
                st.session_state.prd_draft = prd
                qm.record("PRD生成", qm.estimate_tokens(prd))
                prd_path = workspace_dir / "PRD.md"
                prd_path.write_text(prd, encoding="utf-8")
                st.success(f"✅ PRD 已保存到 `{prd_path.absolute()}`")
                if "PRD生成" not in st.session_state.completed_stages:
                    st.session_state.completed_stages.append("PRD生成")
                autosave()

    if st.session_state.get("prd_draft"):
        with st.expander("📄 PRD 当前版本", expanded=True):
            st.markdown(st.session_state.prd_draft)
        if st.button("💾 重新保存 PRD 到 workspace/", key="resave_prd"):
            prd_path = Path(__file__).resolve().parent.parent / "workspace" / "PRD.md"
            prd_path.write_text(st.session_state.prd_draft, encoding="utf-8")
            st.success(f"✅ 已覆盖保存到 `{prd_path.absolute()}`")

    # === 步骤C：生成 CLAUDE.md（LLM 理解 prepare_claude/ 后生成）===
    if st.session_state.get("modeling_plan"):
        if st.button("🤖 生成 CLAUDE.md 到 workspace/", key="gen_claude_md"):
            workspace_dir = Path(__file__).resolve().parent.parent / "workspace"
            prep_dir = workspace_dir / "prepare_claude"
            prep_dir.mkdir(parents=True, exist_ok=True)

            # 确保 prepare_claude/ 下有最新内容
            (prep_dir / "赛题.txt").write_text(st.session_state.get('selected_topic', '')[:2000], encoding="utf-8")
            (prep_dir / "建模方案.md").write_text(st.session_state.get('modeling_plan', '')[:2000], encoding="utf-8")
            prd_text = st.session_state.get('prd_draft', '') or (workspace_dir / "PRD.md").read_text(encoding="utf-8") if (workspace_dir / "PRD.md").exists() else ""
            (prep_dir / "PRD摘要.md").write_text(prd_text[:2000], encoding="utf-8")
            pressure_text = st.session_state.get('pressure_report', '') or ""
            (prep_dir / "压力测试报告.md").write_text(pressure_text[:2000], encoding="utf-8")

            # LLM 读取 prepare_claude/ 下所有文件
            prep_contents = []
            for f in sorted(prep_dir.glob("*")):
                if f.is_file():
                    try:
                        content = f.read_text(encoding="utf-8")[:2000]
                        prep_contents.append(f"### {f.name}\n{content}")
                    except Exception:
                        pass
            prep_text = "\n\n".join(prep_contents)

            with st.spinner("LLM 正在理解 prepare_claude/ 内容并生成 CLAUDE.md..."):
                claude_prompt = f"""你是数学建模竞赛技术经理。请先理解以下全部项目文件内容，然后生成一份 CLAUDE.md 操作手册。

## 项目文件内容
{prep_text}

## CLAUDE.md 结构要求
按以下结构生成 Markdown：

# 数学建模任务 — APMCM Agent 生成
**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M')}

## 1. 任务概述
用自己的话概括赛题和核心目标

## 2. 建模方案摘要
提取方案中的关键方法和数学模型

## 3. 执行计划（表格）
| 阶段 | 任务 | 预期产出 | 验收标准 |
|------|------|----------|----------|

## 4. 技术约束
- 编程语言/库/工具
- 数据格式要求
- 输出规范

## 5. 文件结构
```
workspace/
├── coding/     # 代码输出
├── writing/    # 论文各节
├── picture/    # 图表输出
└── ...
```

## 6. 执行规范
- 编码前: think（设计决策审查）
- 编码中: tdd（测试驱动开发）
- 编码后: check（正确性/健壮性检查）
- 遇Bug: hunt（定位修复）

## 7. 关键风险与注意事项
从压力测试报告中提取

直接输出 Markdown，不要前后缀。"""
                claude_md = gpt_with_retry(claude_prompt, max_tokens=3000)

            claude_path = workspace_dir / "CLAUDE.md"
            claude_path.write_text(claude_md, encoding="utf-8")
            qm.record("CLAUDE.md生成", qm.estimate_tokens(claude_md))
            st.success(f"✅ CLAUDE.md 已由 LLM 理解生成 → `{claude_path.absolute()}`")
            with st.expander("📄 预览 CLAUDE.md"):
                st.markdown(claude_md[:2000])
            autosave()

    # === 步骤D：grill-me 需求对齐循环 ===
    if st.session_state.get("prd_draft"):
        st.subheader("🔥 需求对齐（grill-me）")
        st.caption(f"已对齐 {st.session_state.grill_rounds} 轮 | 建议 2-3 轮后确认最终 PRD")

        user_prd_feedback = st.text_area(
            "对 PRD 有什么不满意或想调整的地方？（直接说，AI 会追问并修订）",
            key="prd_feedback",
            placeholder="例如：我觉得用神经网络更合适 / 时间计划太紧 / 第3章结构不清晰..."
        )
        col_grill, col_confirm = st.columns(2)

        with col_grill:
            if st.button("💬 发起一轮对齐", key="run_grill") and user_prd_feedback:
                with st.spinner("grill-me 追问分析中..."):
                    grill_result = run_grill_me(st.session_state.prd_draft, user_prd_feedback)
                    qm = st.session_state.quota_monitor
                    qm.record("需求对齐", qm.estimate_tokens(grill_result))
                    st.session_state.grill_rounds += 1

                st.subheader("🎯 对齐分析与修订建议")
                st.markdown(grill_result)
                if "需求对齐" not in st.session_state.completed_stages:
                    st.session_state.completed_stages.append("需求对齐")

        with col_confirm:
            if st.button("✅ PRD 已对齐，确认最终版", key="confirm_prd"):
                workspace_dir = Path(__file__).resolve().parent.parent / "workspace"
                workspace_dir.mkdir(parents=True, exist_ok=True)

                st.session_state.prd_final = st.session_state.prd_draft
                (workspace_dir / "PRD.md").write_text(st.session_state.prd_final, encoding="utf-8")
                st.success(f"✅ PRD 已确认并保存到 `{workspace_dir.absolute() / 'PRD.md'}`")
                qm.record("PRD确认", 200)

                if not st.session_state.reference_loaded:
                    from rag import RAG
                    if "rag" in st.session_state and st.session_state.rag.load_references():
                        st.info("📚 reference/ 知识库已加载")
                        st.session_state.reference_loaded = True

                st.session_state.phase = "coding"
                st.session_state.memory_logger.new_stage(st.session_state.phase)
                st.session_state.memory_logger.log_system_event(
                    "进入阶段: coding",
                    f"session_id={st.session_state.session_id}"
                )
                autosave()
                st.rerun()

    st.divider()
    # 综合外部 md 文件重新生成方案
    st.subheader("📥 综合外部参考重新生成方案")
    web_ai_dir = Path(__file__).resolve().parent.parent / "inbox" / "web_ai"
    web_ai_files = []
    if web_ai_dir.exists():
        for f in web_ai_dir.glob("*.md"):
            web_ai_files.append(f.name)
        for f in web_ai_dir.glob("*.txt"):
            web_ai_files.append(f.name)

    # 方式A：从 inbox/web_ai/ 选择已有文件
    extra_context = ""
    if web_ai_files:
        selected_ref = st.multiselect("从 inbox/web_ai/ 选择参考文件", web_ai_files, key="ref_files")
        if selected_ref:
            for fn in selected_ref:
                extra_context += f"\n\n### {fn}\n{(web_ai_dir / fn).read_text(encoding='utf-8')[:3000]}"

    # 方式B：直接上传 .md 文件
    uploaded_md = st.file_uploader("或上传 .md 文件", type=["md", "txt"], key="upload_md_ref")
    if uploaded_md:
        extra_context += f"\n\n### {uploaded_md.name}\n{uploaded_md.read().decode('utf-8')[:3000]}"

    user_approach = st.text_area("补充建模方向建议(可选)", key="user_approach_input",
                                  placeholder="例如:希望使用动态规划方法...")
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("🔄 重新生成方案"):
            rubbish_downstream("modeling")
            if user_approach:
                st.session_state.user_approach = user_approach
            # 保存外部参考上下文供下次生成使用
            if extra_context:
                st.session_state._modeling_refs = extra_context
            st.session_state.pop("modeling_plan", None)
            st.rerun()
    with col2:
        if st.button("🧪 压力测试", type="primary"):
            logger.log_user_feedback("modeling", "用户确认建模方案,进入压力测试")
            st.session_state.phase = "pressure_test"
            st.session_state.memory_logger.new_stage(st.session_state.phase)
            st.session_state.memory_logger.log_system_event(
                "进入阶段: pressure_test",
                f"session_id={st.session_state.session_id}"
            )
            autosave()
            st.rerun()
    with col3:
        if st.button("⏭️ 跳过测试"):
            logger.log_user_feedback("modeling", "用户跳过测试,直接进入代码生成")
            st.session_state.phase = "coding"
            st.session_state.memory_logger.new_stage(st.session_state.phase)
            st.session_state.memory_logger.log_system_event(
                "进入阶段: coding",
                f"session_id={st.session_state.session_id}"
            )
            autosave()
            st.rerun()

# ============ Phase 4: Pressure Test ============
elif st.session_state.phase == "pressure_test":
    st.header("🧪 建模方案压力测试")

    modeling_plan = st.session_state.modeling_plan

    if "pressure_test_result" not in st.session_state:
        with st.spinner("正在进行压力测试..."):
            prompt = get_pressure_test_prompt(modeling_plan)
            test_result = gpt_with_retry(prompt)
            st.session_state.pressure_test_result = test_result
            logger.log_pressure_test(test_result)

    st.markdown(st.session_state.pressure_test_result)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 返回修改方案"):
            st.session_state.phase = "modeling"
            st.session_state.memory_logger.new_stage(st.session_state.phase)
            st.session_state.memory_logger.log_system_event(
                "进入阶段: modeling",
                f"session_id={st.session_state.session_id}"
            )
            autosave()
            st.rerun()
    with col2:
        if st.button("✅ 方案通过,对齐用户期望", type="primary"):
            st.session_state.phase = "grill_me"
            st.session_state.memory_logger.new_stage(st.session_state.phase)
            st.session_state.memory_logger.log_system_event(
                "进入阶段: grill_me",
                f"session_id={st.session_state.session_id}"
            )
            autosave()
            st.rerun()

# ============ Phase 5: Grill Me ============
elif st.session_state.phase == "grill_me":
    st.header("🎯 用户期望对齐")

    modeling_plan = st.session_state.modeling_plan
    user_expectation = st.text_area(
        "请输入你的预期和目标",
        value=st.session_state.get("user_expectation", ""),
        height=150,
        placeholder="例如:我期望方案能够获得省级一等奖,并且代码实现简单..."
    )

    if st.button("🔍 评估方案是否符合预期", type="primary") and user_expectation:
        with st.spinner("正在评估..."):
            st.session_state.user_expectation = user_expectation
            prompt = get_grill_me_prompt(modeling_plan, user_expectation)
            grill_result = gpt_with_retry(prompt)
            st.session_state.grill_result = grill_result
            logger.log_grill_me(grill_result)

    if "grill_result" in st.session_state:
        st.markdown(st.session_state.grill_result)
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 返回调整方案"):
                st.session_state.phase = "modeling"
                st.session_state.memory_logger.new_stage(st.session_state.phase)
                st.session_state.memory_logger.log_system_event(
                    "进入阶段: modeling",
                    f"session_id={st.session_state.session_id}"
                )
                autosave()
                st.rerun()
        with col2:
            if st.button("✅ 方案对齐,开始编码", type="primary"):
                st.session_state.phase = "coding"
                st.session_state.memory_logger.new_stage(st.session_state.phase)
                st.session_state.memory_logger.log_system_event(
                    "进入阶段: coding",
                    f"session_id={st.session_state.session_id}"
                )
                autosave()
                st.rerun()

# ============ Phase 6: Coding ============
elif st.session_state.phase == "coding":
    st.header("💻 代码生成")
    WS = Path(__file__).resolve().parent.parent / "workspace"
    n = st.session_state.get("selected_topic_idx", 0) + 1

    if st.button("⬅️ 返回建模方案", key="back_to_modeling"):
        st.session_state.phase = "modeling"
        st.session_state.memory_logger.new_stage("modeling")
        autosave()
        st.rerun()
    has_figure = st.session_state.get("figure_descriptions") is not None
    if st.button("→ 图表方案", key="next_to_figure", disabled=not has_figure,
                 help="需先生成图表方案" if not has_figure else "进入图表方案阶段"):
        st.session_state.phase = "figure"
        st.session_state.memory_logger.new_stage("figure")
        st.session_state.memory_logger.log_system_event("进入阶段: figure")
        autosave()
        st.rerun()

    # ── 准备文件夹：保存所有输入文件为 .md ──
    prepare_dir = WS / "coding" / str(n) / f"prepare_{n}"
    prepare_dir.mkdir(parents=True, exist_ok=True)

    # 保存题目
    topic_text = st.session_state.get("selected_topic", "")
    if topic_text:
        (prepare_dir / "topic.md").write_text(str(topic_text)[:3000], encoding="utf-8")

    # 保存建模方案
    plan = st.session_state.get("modeling_plan", "")
    if plan:
        (prepare_dir / "modeling_plan.md").write_text(str(plan)[:3000], encoding="utf-8")

    # 保存相似题/论文参考
    sims = st.session_state.get("selected_sims", {})
    if sims:
        ref_text = "# 相似赛题\n\n"
        for q in sims.get("sim_questions", [])[:3]:
            ref_text += f"- {q[:500]}\n\n"
        ref_text += "# 获奖论文\n\n"
        for p in sims.get("sim_papers", [])[:3]:
            ref_text += f"- {p[:500]}\n\n"
        (prepare_dir / "references.md").write_text(ref_text, encoding="utf-8")

    # 复制 PRD.md
    prd_src = WS / "PRD.md"
    if prd_src.exists():
        (prepare_dir / "PRD.md").write_text(prd_src.read_text(encoding="utf-8")[:3000], encoding="utf-8")

    # 复制 CLAUDE.md
    claude_src = WS / "CLAUDE.md"
    if claude_src.exists():
        (prepare_dir / "CLAUDE.md").write_text(claude_src.read_text(encoding="utf-8")[:3000], encoding="utf-8")

    # 复制数据文件
    data_dir = WS.parent / "inbox" / "data"
    if data_dir.exists():
        for df in data_dir.glob("*"):
            if df.suffix.lower() in (".csv", ".xlsx", ".xls"):
                shutil.copy2(str(df), str(prepare_dir / df.name))

    # 展示给用户
    st.subheader(f"📁 选题 {n} 的准备工作")
    st.caption(f"以下文件将从 `{prepare_dir.absolute()}` 读取用于代码生成：")
    prep_files = sorted(prepare_dir.glob("*"))
    if prep_files:
        for f in prep_files:
            if f.is_file():
                st.text(f"  📄 {f.name}")
    else:
        st.warning("prepare 文件夹为空，请先在建模阶段生成 PRD.md 和 CLAUDE.md")

    # 代码输出目录
    code_output_dir = WS / "coding" / str(n)
    code_output_dir.mkdir(parents=True, exist_ok=True)

    st.markdown("---")

    # 生成按钮
    if st.button("🚀 开始生成代码", type="primary", key="btn_gen_code"):
        with st.spinner("正在从 prepare 文件夹读取所有 .md 文件并生成代码..."):
            prompt_parts = []
            for f in sorted(prepare_dir.glob("*.md")):
                content = f.read_text(encoding="utf-8")[:2000]
                prompt_parts.append(f"## {f.name}\n{content}")
            prompt = "\n\n".join(prompt_parts)
            prompt = f"请根据以下全部文件内容编写完整的数学建模 Python 代码：\n\n{prompt}"
            prompt += f"\n\n保存代码到 {code_output_dir.absolute()}，图表到 {WS.absolute() / 'picture'}。"
            coding_result = gpt_with_retry(prompt, max_tokens=8000)
            st.session_state.coding_result = coding_result
            logger.log_coding(coding_result)
            # 保存代码
            (code_output_dir / "solution.py").write_text(coding_result, encoding="utf-8")
            st.success(f"代码已生成并保存到 `{code_output_dir.absolute()}`")
            autosave()

    # 代码显示
    if st.session_state.get("coding_result"):
        st.code(st.session_state.coding_result, language="python")

        # 代码审查
        with st.expander("🔍 代码三重审查（think + check + TDD）", expanded=False):
            if st.button("运行 Skill 审查", key="run_code_check"):
                with st.spinner("运行 think / check / tdd 审查..."):
                    review = run_code_check(st.session_state.coding_result)
                    qm.record("代码审查", qm.estimate_tokens(review))
                    st.markdown(review)

        # ── 论文节：LLM 先理解再总结到 prepare 文件夹 ──
        st.subheader("📝 同步生成对应论文节")
        writing_prep = WS / "writing" / str(n) / "prepare"
        writing_prep.mkdir(parents=True, exist_ok=True)

        # LLM 理解代码并总结
        if not (writing_prep / "code_summary.md").exists():
            if st.button("🧠 理解代码并生成摘要", key="btn_code_summary"):
                with st.spinner("LLM 正在阅读代码..."):
                    from model import gpt_with_retry
                    code_summary = gpt_with_retry(
                        f"请阅读以下 Python 代码，用中文总结：1)核心算法 2)输入输出 3)关键参数 4)主要函数。\n\n{st.session_state.coding_result[:3000]}",
                        max_tokens=800)
                    (writing_prep / "code_summary.md").write_text(code_summary, encoding="utf-8")
                st.success("代码摘要已保存")
                st.rerun()

        # LLM 理解建模方案并总结
        if not (writing_prep / "model_summary.md").exists():
            if st.button("🧠 理解建模方案并生成摘要", key="btn_model_summary"):
                with st.spinner("LLM 正在阅读建模方案..."):
                    from model import gpt_with_retry
                    model_summary = gpt_with_retry(
                        f"请阅读以下建模方案，用中文总结：1)数学模型 2)假设条件 3)求解方法 4)创新点。\n\n{st.session_state.get('modeling_plan','')[:3000]}",
                        max_tokens=800)
                    (writing_prep / "model_summary.md").write_text(model_summary, encoding="utf-8")
                st.success("建模方案摘要已保存")
                st.rerun()

        # 展示 prepare 文件夹已有哪些文件
        existing_preps = list(writing_prep.glob("*.md"))
        if existing_preps:
            st.caption(f"prepare 文件夹 ({writing_prep.absolute()}):")
            for f in existing_preps:
                st.text(f"  ✅ {f.name}")

        # 生成论文节
        if existing_preps:
            section_name = st.selectbox("选择论文章节",
                                         ["3.1 模型建立", "3.2 求解方法", "4.1 结果分析", "4.2 敏感性分析"])
            if st.button("📖 生成这一节论文初稿", key="gen_paper_section"):
                with st.spinner(f"从 prepare 文件夹读取并生成 {section_name}..."):
                    from model import gpt_with_retry
                    prep_text = ""
                    for f in sorted(writing_prep.glob("*.md")):
                        prep_text += f"\n## {f.stem}\n{f.read_text(encoding='utf-8')[:1500]}"
                    section = gpt_with_retry(
                        f"根据以下项目摘要，写出数学建模论文的 {section_name} 节。学术语言，含 LaTeX 公式，300-500字。\n\n{prep_text}",
                        max_tokens=800)
                    qm.record(f"论文节-{section_name}", qm.estimate_tokens(section))
                    st.session_state.paper_sections[section_name] = section
                    safe_name = section_name.replace(" ", "_").replace(".", "")
                    (WS / "writing" / str(n) / f"{safe_name}.md").write_text(section, encoding="utf-8")
                st.markdown(section)
                st.success(f"✅ {section_name} 已保存")

        # 导航
        st.info("👆 确认代码后进入下一阶段")
        col_ok, col_redo = st.columns(2)
        with col_ok:
            if st.button("✅ 符合预期，继续", key="coding_ok"):
                if "代码生成" not in st.session_state.completed_stages:
                    st.session_state.completed_stages.append("代码生成")
                st.session_state.phase = "figure"
                st.session_state.memory_logger.new_stage(st.session_state.phase)
                st.session_state.memory_logger.log_system_event("进入阶段: figure", f"session_id={st.session_state.session_id}")
                autosave()
                st.rerun()
        with col_redo:
            if st.button("🔄 重新生成", key="coding_redo"):
                rubbish_downstream("coding")
                st.session_state.pop("coding_result", None)
                st.rerun()

# ============ Phase 7: Figure ============
elif st.session_state.phase == "figure":
    st.header("📊 图表生成方案")

    if st.button("⬅️ 返回代码生成", key="back_to_coding"):
        st.session_state.phase = "coding"
        st.session_state.memory_logger.new_stage("coding")
        autosave()
        st.rerun()
    has_paper = st.session_state.get("paper_draft") is not None
    if st.button("→ 论文初稿", key="next_to_paper", disabled=not has_paper,
                 help="需先生成论文初稿" if not has_paper else "进入论文初稿阶段"):
        st.session_state.phase = "paper"
        st.session_state.memory_logger.new_stage("paper")
        st.session_state.memory_logger.log_system_event("进入阶段: paper")
        autosave()
        st.rerun()

    if st.session_state.figure_descriptions is None:
        with st.spinner("正在设计图表方案..."):
            prompt = get_figure_prompt(
                st.session_state.selected_topic,
                st.session_state.modeling_plan,
                st.session_state.coding_result or ""
            )
            prompt += "\n\n每个图表请标注好题号和名称，格式如：**图表1-1 线性回归曲线**"
            figure_desc = gpt_with_retry(prompt)
            st.session_state.figure_descriptions = figure_desc
            logger.log_figure(figure_desc)
            pic_dir = Path(__file__).resolve().parent.parent / "workspace" / "picture"
            pic_dir.mkdir(parents=True, exist_ok=True)
            (pic_dir / "figure_plan.md").write_text(figure_desc, encoding="utf-8")

    st.markdown(st.session_state.figure_descriptions)

    st.info("💡 安装 scipilot-figure-skill 后可直接在Python中生成图表:")
    st.code(
        "from skills_bridge import profile_data, setup_style, export_figure\n"
        "# 示例: setup_style('nature'); export_figure(fig, 'result')",
        language="python"
    )

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 重新生成图表方案"):
            rubbish_downstream("figure")
            st.session_state.pop("figure_descriptions", None)
            st.rerun()
    with col2:
        if st.button("📝 生成论文初稿", type="primary"):
            st.session_state.phase = "paper"
            st.session_state.memory_logger.new_stage(st.session_state.phase)
            st.session_state.memory_logger.log_system_event(
                "进入阶段: paper",
                f"session_id={st.session_state.session_id}"
            )
            autosave()
            st.rerun()

# ============ Phase 8: Paper ============
elif st.session_state.phase == "paper":
    st.header("📝 论文初稿")
    WS = Path(__file__).resolve().parent.parent / "workspace"
    n = st.session_state.get("selected_topic_idx", 0) + 1
    writing_dir = WS / "writing" / str(n)

    if st.button("⬅️ 返回图表方案", key="back_to_figure"):
        st.session_state.phase = "figure"
        st.session_state.memory_logger.new_stage("figure")
        autosave()
        st.rerun()
    has_polish = st.session_state.get("polished_paper") is not None
    if st.button("→ 论文润色", key="next_to_polish", disabled=not has_polish,
                 help="需先生成润色版论文" if not has_polish else "进入论文润色阶段"):
        st.session_state.phase = "polish"
        st.session_state.memory_logger.new_stage("polish")
        st.session_state.memory_logger.log_system_event("进入阶段: polish")
        autosave()
        st.rerun()

    # 汇总各阶段论文节
    section_files = list((writing_dir).glob("*.md")) if writing_dir.exists() else []
    st.caption(f"已有 {len(section_files)} 节论文初稿在 `{writing_dir.absolute()}`")
    if section_files:
        for f in sorted(section_files):
            st.text(f"  📄 {f.name}")

    if st.session_state.paper_draft is None:
        with st.spinner("正在综合各节生成完整论文..."):
            # 优先从写作文件夹汇总，否则用原始 prompt
            if section_files:
                parts = []
                for f in sorted(section_files):
                    parts.append(f.read_text(encoding="utf-8"))
                section_text = "\n\n".join(parts)
                prompt = f"""请将以下各节论文内容整合为一篇完整的数学建模竞赛论文。
要求：摘要→问题重述→模型建立→求解→检验→评价→参考文献，学术语言，LaTeX 公式。

{section_text}"""
            else:
                prompt = get_paper_writing_prompt(
                    st.session_state.selected_topic,
                    st.session_state.modeling_plan,
                    st.session_state.coding_result or "",
                    st.session_state.figure_descriptions or ""
                )
            paper_draft = gpt_with_retry(prompt, max_tokens=8000)
            st.session_state.paper_draft = paper_draft
            logger.log_paper_draft(paper_draft)
            writing_dir.mkdir(parents=True, exist_ok=True)
            (writing_dir / "paper_complete.md").write_text(paper_draft, encoding="utf-8")
            st.caption(f"完整论文已保存到 `{writing_dir.absolute() / 'paper_complete.md'}`")

    st.markdown(st.session_state.paper_draft)

    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("🔄 重新生成论文"):
            rubbish_downstream("paper")
            st.session_state.pop("paper_draft", None)
            st.rerun()
    with col2:
        if st.button("✨ 润色论文", type="primary"):
            st.session_state.phase = "polish"
            st.session_state.memory_logger.new_stage(st.session_state.phase)
            st.session_state.memory_logger.log_system_event(
                "进入阶段: polish",
                f"session_id={st.session_state.session_id}"
            )
            autosave()
            st.rerun()
    with col3:
        if st.button("📥 导出完成"):
            st.session_state.phase = "done"
            st.session_state.memory_logger.new_stage(st.session_state.phase)
            st.session_state.memory_logger.log_system_event(
                "进入阶段: done",
                f"session_id={st.session_state.session_id}"
            )
            autosave()
            st.rerun()

# ============ Phase 9: Polish ============
elif st.session_state.phase == "polish":
    st.header("✨ 论文润色")

    if st.button("⬅️ 返回论文初稿", key="back_to_paper"):
        st.session_state.phase = "paper"
        st.session_state.memory_logger.new_stage("paper")
        autosave()
        st.rerun()
    # done is always reachable from polish
    if st.button("→ 完成", key="next_to_done"):
        st.session_state.phase = "done"
        st.session_state.memory_logger.new_stage("done")
        st.session_state.memory_logger.log_system_event("进入阶段: done")
        autosave()
        st.rerun()

    polish_type = st.selectbox("润色类型", ["润色", "翻译为英文", "学术语法修正", "逻辑优化"])

    if st.button("🚀 开始润色", type="primary"):
        with st.spinner(f"正在进行{polish_type}..."):
            prompt = get_polish_prompt(st.session_state.paper_draft, polish_type)
            polished = gpt_with_retry(prompt, max_tokens=8000)
            st.session_state.polished_paper = polished
            logger.log_paper_polish(polished)
            writing_dir = Path(__file__).resolve().parent.parent / "workspace" / "writing"
            writing_dir.mkdir(parents=True, exist_ok=True)
            (writing_dir / "paper_polished.md").write_text(polished, encoding="utf-8")

    if "polished_paper" in st.session_state:
        st.markdown("### 润色结果")
        st.markdown(st.session_state.polished_paper)

        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 重新润色"):
                st.session_state.pop("polished_paper", None)
                st.rerun()
        with col2:
            if st.button("📥 完成,查看总结", type="primary"):
                st.session_state.phase = "done"
                st.session_state.memory_logger.new_stage(st.session_state.phase)
                st.session_state.memory_logger.log_system_event(
                    "进入阶段: done",
                    f"session_id={st.session_state.session_id}"
                )
                autosave()
                st.rerun()

        # === 论文导出 ===
        st.divider()
        st.subheader("📄 导出论文")
        export_col1, export_col2 = st.columns(2)

        with export_col1:
            if st.button("📥 导出为 Word 文档 (.docx)", key="export_docx"):
                try:
                    from docx import Document as DocxDocument
                    from docx.shared import Pt, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH

                    paper_text = st.session_state.get(
                        "polished_paper",
                        st.session_state.get("paper_draft", "")
                    )

                    doc = DocxDocument()

                    title = doc.add_heading("数学建模竞赛论文", 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    for line in paper_text.split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        if line.startswith("# "):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith("## "):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith("### "):
                            doc.add_heading(line[4:], level=3)
                        elif line.startswith("$$") or line.startswith("\\begin"):
                            p = doc.add_paragraph()
                            run = p.add_run(f"[公式] {line}")
                            run.font.name = "Courier New"
                            run.font.size = Pt(9)
                        else:
                            doc.add_paragraph(line)

                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)

                    st.download_button(
                        label="⬇️ 点击下载 .docx",
                        data=buf.getvalue(),
                        file_name=f"apmcm_paper_{st.session_state.session_id}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_docx"
                    )
                except ImportError:
                    st.error("请先安装 python-docx：pip install python-docx")

        with export_col2:
            if st.button("📝 导出为 Markdown 文件", key="export_md"):
                paper_text = st.session_state.get(
                    "polished_paper",
                    st.session_state.get("paper_draft", "")
                )
                workspace = Path(__file__).resolve().parent.parent / "workspace" / st.session_state.session_id
                workspace.mkdir(parents=True, exist_ok=True)
                paper_path = workspace / "paper.md"
                paper_path.write_text(paper_text, encoding="utf-8")
                st.success(f"✅ 论文已保存到 `{paper_path.absolute()}`")
                st.code(
                    f"# 转换为 PDF（需要安装 pandoc 和 xelatex）\n"
                    f"cd {workspace.absolute()}\n"
                    f"pandoc paper.md -o paper.pdf --pdf-engine=xelatex "
                    f"-V mainfont='SimSun' -V geometry:margin=2.5cm",
                    language="bash"
                )

# ============ Done ============
elif st.session_state.phase == "done":
    st.header("🎉 工作流程完成!")

    summary = logger.get_summary()
    st.markdown(f"""
    ### 会话总结
    - **Session ID**: {summary['session_id']}
    - **总步骤数**: {summary['total_steps']}
    - **日志文件**: `{summary['log_file']}`
    - **Markdown日志**: `{summary['md_file']}`
    """)

    st.markdown("### 各阶段完成情况")
    for phase, count in summary['phases'].items():
        st.markdown(f"- {phase}: {count} 步")

    st.divider()
    st.markdown("### 📄 最终产出")

    tab1, tab2, tab3, tab4 = st.tabs(["建模方案", "代码", "图表方案", "论文"])

    with tab1:
        if st.session_state.modeling_plan:
            st.markdown(st.session_state.modeling_plan)
    with tab2:
        if st.session_state.coding_result:
            st.code(st.session_state.coding_result, language="python")
    with tab3:
        if st.session_state.figure_descriptions:
            st.markdown(st.session_state.figure_descriptions)
    with tab4:
        paper = st.session_state.get("polished_paper") or st.session_state.paper_draft
        if paper:
            st.markdown(paper)

    if st.button("🔄 开始新会话"):
        for key in list(st.session_state.keys()):
            if key not in ["rag"]:
                del st.session_state[key]
        st.session_state.logger = WorkflowLogger()
        st.session_state.rag = rag
        st.session_state.phase = "input"
        st.session_state.session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        st.query_params["session"] = st.session_state.session_id
        st.session_state.memory_logger = MemoryLogger(st.session_state.session_id)
        st.session_state.memory_logger.new_stage("input")
        st.session_state.memory_logger.log_system_event(
            "新会话开始",
            f"session_id={st.session_state.session_id}"
        )
        st.rerun()

# ═══════════════════════════════════════════════════════════
# 💬 对话面板（始终可见，与工作区独立）
# ═══════════════════════════════════════════════════════════
st.divider()
chat_col, chat_info_col = st.columns([4, 1])
with chat_col:
    st.header("💬 与 Agent 对话")
with chat_info_col:
    chat_count = len(st.session_state.get("chat_history", []))
    st.caption(f"共 {chat_count} 条消息")

chat_container = st.container(height=400, border=True)
with chat_container:
    if st.session_state.get("chat_history"):
        for msg in st.session_state.chat_history:
            role_label = "🧑 你" if msg["role"] == "user" else "🤖 Agent"
            with st.chat_message(msg["role"]):
                st.markdown(f"**{role_label}** — {msg['content'][:2000]}")
    else:
        st.caption("对话记录将在这里显示。上传赛题后即可与 Agent 交流。")

st.caption("")
chat_input = st.chat_input("向Agent提问或提供反馈...")
if chat_input:
    webai_content = st.session_state.get("webai_imported_content", "")
    if webai_content and ("使用网页AI方案" in chat_input or "应用网页AI回复" in chat_input):
        chat_input = f"【网页版AI方案】\n{webai_content}\n\n【用户指令】\n{chat_input}"
        st.session_state.webai_imported_content = ""

    if ("分析数据" in chat_input or "数据文件" in chat_input) \
            and st.session_state.get("loaded_data_files"):
        data_inject = "\n".join(
            f"【{k}】\n{v}" for k, v in st.session_state.loaded_data_files.items()
        )
        chat_input = f"以下是已加载的数据文件内容：\n{data_inject}\n\n用户问题：{chat_input}"

    st.session_state.chat_history.append({"role": "user", "content": chat_input})
    st.session_state.memory_logger.log_message("user", chat_input)
    with st.spinner("Agent思考中..."):
        ctx = f"""当前阶段: {st.session_state.phase}
已选赛题: {(st.session_state.get('selected_topic') or '未选择')[:500]}
建模方案摘要: {(st.session_state.get('modeling_plan') or '未生成')[:800]}
压力测试: {(st.session_state.get('pressure_report') or '未生成')[:300]}
PRD: {(st.session_state.get('prd_final') or st.session_state.get('prd_draft') or '未生成')[:300]}
代码摘要: {(st.session_state.get('coding_result') or '未生成')[:500]}
论文摘要: {(st.session_state.get('paper_draft') or '未生成')[:300]}"""
        messages_to_send = []
        if st.session_state.get("context_summary"):
            messages_to_send.append({
                "role": "system",
                "content": st.session_state.context_summary
            })
            st.session_state.context_summary = ""
        messages_to_send.append({
            "role": "system",
            "content": f"""你是APMCM数学建模Agent助手，全程参与用户的建模竞赛准备。你可以：
- 解释当前阶段的决策逻辑和方法选择理由
- 回答关于建模方案、代码、论文的任何问题
- 根据用户反馈调整方向和策略
请用中文回答，简洁有针对性。

{ctx}""",
        })
        recent_history = st.session_state.chat_history[-20:]
        messages_to_send.extend(recent_history)
        response = chat(messages_to_send)
    st.session_state.chat_history.append({"role": "assistant", "content": response})
    st.session_state.memory_logger.log_message("assistant", response)
    st.rerun()

autosave()
